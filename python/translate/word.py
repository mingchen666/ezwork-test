import threading
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn
import translate
import common
import os
import sys
import time
import datetime
import zipfile
import xml.etree.ElementTree as ET
import rediscon

def start(trans):
    # 允许的最大线程
    threads=trans['threads']
    if threads is None or threads=="" or int(threads)<0:
        max_threads=10
    else:
        max_threads=int(threads)
    # 当前执行的索引位置
    run_index=0
    max_chars=1000
    start_time = datetime.datetime.now()
    # 创建Document对象，加载Word文件
    try:
        document = Document(trans['file_path'])
    except Exception as e:
        translate.error(trans['id'], "无法访问该文档")
        return False
    texts=[]
    api_url=trans['api_url']
    trans_type=trans['type']
    if trans_type=="trans_text_only_inherit":
        # 仅文字-保留原文-继承原版面
        read_rune_text(document, texts)
    elif trans_type=="trans_text_only_new" or trans_type=="trans_text_both_new":
        # 仅文字-保留原文-重排
        read_paragraph_text(document, texts)
    elif trans_type=="trans_text_both_inherit":
        # 仅文字-保留原文-重排/继承原版面
        read_rune_text(document, texts)
    elif trans_type=="trans_all_only_new":
        # 全部内容-仅译文-重排版面
        read_paragraph_text(document, texts)
    elif trans_type=="trans_all_only_inherit":
        # 全部内容-仅译文-重排版面/继承原版面
        read_rune_text(document, texts)
    elif trans_type=="trans_all_both_new":
        # 全部内容-保留原文-重排版面
        read_paragraph_text(document, texts)
    elif trans_type=="trans_all_both_inherit":
        # 全部内容-保留原文-继承原版面
        read_rune_text(document, texts)

    read_comments_from_docx(trans['file_path'], texts)
    read_insstd_from_docx(trans['file_path'], texts)
    #print(texts)
    #exit()
    max_run=max_threads if len(texts)>max_threads else len(texts)
    #mredis=rediscon.get_conn()
    #threading_num=int(mredis.get(api_url))
    #mredis.set(api_url,threading_num+max_run)
    event=threading.Event()
    before_active_count=threading.activeCount()
    while run_index<=len(texts)-1:
        if threading.activeCount()<max_run+before_active_count:
            if not event.is_set():
                thread = threading.Thread(target=translate.get,args=(trans,event,texts,run_index))
                thread.start()
                run_index+=1
            else:
                return False
    
    while True:
        if event.is_set():
            return False
        complete=True
        for text in texts:
            if not text['complete']:
                complete=False
        if complete:
            break
        else:
            time.sleep(1)
    # print(texts)
    # print("翻译文本-结束")

    text_count=0
    if trans_type=="trans_text_only_inherit":
        # 仅文字-仅译文-继承原版面。
        write_only_new(document, texts, text_count, True) # DONE
    elif trans_type=="trans_text_only_new":
        # 仅文字-仅译文-重排
        write_paragraph_text(document, texts, text_count, True) #DONE
    elif trans_type=="trans_text_both_new":
        # 仅文字-保留原文-重排
        write_both_new(document, texts, text_count, True) #DONE
    elif trans_type=="trans_text_both_inherit":
        # 仅文字-保留原文-继承原版面
        write_rune_both(document, texts, text_count, True) #DONE
    elif trans_type=="trans_all_only_new":
        # 全部内容-仅译文-重排版面
        write_paragraph_text(document, texts, text_count, False) #DONE
    elif trans_type=="trans_all_only_inherit":
        # 全部内容-仅译文-继承原版面
        write_only_new(document, texts, text_count, False) #DONE
    elif trans_type=="trans_all_both_new":
        # 全部内容-保留原文-重排版面
        write_both_new(document, texts, text_count, False) #DONE
    elif trans_type=="trans_all_both_inherit":
        # 全部内容-保留原文-继承原版面
        write_rune_both(document, texts, text_count, False) #DONE

    # print("编辑文档-结束")
    # print(datetime.datetime.now())
    docx_path=trans['target_file']
    document.save(docx_path)
    # 替换批注数据
    modify_comment_in_docx(docx_path, texts)
    modify_inssdt_in_docx(docx_path, texts)
    end_time = datetime.datetime.now()
    spend_time=common.display_spend(start_time, end_time)
    if trans['run_complete']:
        translate.complete(trans,text_count,spend_time)
    return True


def read_paragraph_text(document, texts):
    for paragraph in document.paragraphs:
        append_text(paragraph.text, texts)

    for table in document.tables:
        for row in table.rows:
            start_span=0
            for cell in row.cells:
                read_cell_text(cell, texts)

def write_paragraph_text(document, texts, text_count, onlyText):
    for paragraph in document.paragraphs:
        replace_paragraph_text(paragraph, texts, text_count, onlyText, False)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                write_paragraph_text(cell, texts, text_count, onlyText)

def write_both_new(document, texts, text_count, onlyText):
    for paragraph in document.paragraphs:
        replace_paragraph_text(paragraph, texts, text_count, onlyText, True)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                write_both_new(cell, texts, text_count, onlyText)

def read_cell_text(cell, texts):
    for index,paragraph in enumerate(cell.paragraphs):
        append_text(paragraph.text, texts)

def write_cell_text(cell, texts):
    for index,paragraph in enumerate(cell.paragraphs):
        if check_text(paragraph.text) and len(texts)>0:
            item=texts.pop(0)
            # paragraph.runs[0].text=item.get('text',"")
            for index,run in enumerate(paragraph.runs):
                if index==0:
                    run.text=item.get('text',"")
                else:
                    run.clear()

def read_rune_text(document, texts):
    for paragraph in document.paragraphs:
        line_spacing=paragraph.paragraph_format.line_spacing
        # print("line_spacing:",line_spacing)
        read_run(paragraph.runs, texts)
        # print(line_spacing_unit)
        if len(paragraph.hyperlinks)>0:
            for hyperlink in paragraph.hyperlinks:
                read_run(hyperlink.runs, texts)

    # print("翻译文本--开始")
    # print(datetime.datetime.now())
    for table in document.tables:
        for row in table.rows:
            start_span=0
            for cell in row.cells:
                read_cell_text(cell, texts)
                # start_span+=1
                # # if start_span==cell.grid_span:
                # #     start_span=0
                #     # read_cell(cell, texts)
                # for index,paragraph in enumerate(cell.paragraphs):

                #     read_run(paragraph.runs, texts)

                #     if len(paragraph.hyperlinks)>0:
                #         for hyperlink in paragraph.hyperlinks:
                #             read_run(hyperlink.runs, texts)


def write_only_new(document, texts, text_count, onlyText):
    for paragraph in document.paragraphs:
        text_count+=write_run(paragraph.runs, texts)

        if len(paragraph.hyperlinks)>0:
            for hyperlink in paragraph.hyperlinks:
                text_count+=write_run(hyperlink.runs, texts)

        if onlyText:
            clear_image(paragraph)

    for table in document.tables:
        for row in table.rows:
            start_span=0
            for cell in row.cells:
                write_cell_text(cell, texts)
                # start_span+=1
                # if start_span==cell.grid_span:
                #     start_span=0
                    # text_count+=write_cell(cell, texts)
                # for paragraph in cell.paragraphs:
                #     text_count+=write_run(paragraph.runs, texts)

                #     if len(paragraph.hyperlinks)>0:
                #         for hyperlink in paragraph.hyperlinks:
                #             text_count+=write_run(hyperlink.runs, texts)

#保留原译文
def write_rune_both(document, texts, text_count, onlyText):
    for paragraph in document.paragraphs:
        # print(paragraph.text)
        if(len(paragraph.runs)>0):
            paragraph.runs[-1].add_break()
            add_paragraph_run(paragraph, paragraph.runs, texts, text_count)
        if len(paragraph.hyperlinks)>0:
            for hyperlink in paragraph.hyperlinks:
                hyperlink.runs[-1].add_break()
                add_paragraph_run(paragraph, hyperlink.runs, texts, text_count)
        if onlyText:
            clear_image(paragraph)
       
        # text_count+=write_run(paragraph.runs, texts)
    for table in document.tables:
        for row in table.rows:
            # start_span=0
            for cell in row.cells:
                # start_span+=1
                # if start_span==cell.grid_span:
                #     start_span=0
                    # text_count+=write_cell(cell, texts)
                for paragraph in cell.paragraphs:
                    replace_paragraph_text(paragraph, texts, text_count, onlyText, True)

                    if len(paragraph.hyperlinks)>0:
                        for hyperlink in paragraph.hyperlinks:
                            replace_paragraph_text(hyperlink, texts, text_count, onlyText, True)

def read_run(runs,texts):
    # text=""
    if len(runs)>0 or len(texts)==0:
        for index,run in enumerate(runs):
            append_text(run.text, texts)
        #     if run.text=="":
        #         if len(text)>0 and not common.is_all_punc(text):        
        #             texts.append({"text":text, "complete":False})
        #             text=""
        #     else:
        #         text+=run.text
        # if len(text)>0 and not common.is_all_punc(text):
        #     texts.append({"text":text, "complete":False})

def append_text(text, texts):
    if check_text(text):
        # print(text)
        texts.append({"text":text, "type":"text", "complete":False})

def append_comment(text, comment_id, texts):
    if check_text(text):
        texts.append({"text":text, "type":"comment","comment_id":comment_id, "complete":False})

def check_text(text):
    return text!=None and len(text)>0 and not common.is_all_punc(text) 

def write_run(runs,texts):
    text_count=0
    if len(runs)==0:
        return text_count
    text=""
    for index,run in enumerate(runs):
        text=run.text
        if check_text(text) and len(texts)>0:
            item=texts.pop(0)
            text_count+=item.get('count',0)
            run.text=item.get('text',"")
        
        # if run.text=="":
        #     if len(text)>0 and not common.is_all_punc(text) and len(texts)>0:
        #         item=texts.pop(0)
        #         text_count+=item.get('count',0)
        #         runs[index-1].text=item.get('text',"")
        #         text=""
        # else:
        #     text+=run.text
        #     run.text=""
    # if len(text)>0 and not common.is_all_punc(text) and len(texts)>0:
    #     item=texts.pop(0)
    #     text_count+=item.get('count',0)
    #     runs[0].text=item.get('text',"")
    return text_count


def read_cell(cell,texts):
    append_text(cell.text, texts)


def write_cell(cell,texts):
    text=cell.text
    text_count=0
    if check_text(text) and len(texts)>0:
        item=texts.pop(0)
        text_count+=item.get('count',0)
        cell.text=item.get('text',"")
    return text_count

def add_paragraph_run(paragraph, runs, texts, text_count):
    for index,run in enumerate(runs):
        if check_text(run.text) and len(texts)>0:
            item=texts.pop(0)
            text_count+=item.get('count',0)
            new_run=paragraph.add_run(item.get('text',""), run.style)
            set_run_style(new_run, run)
    set_paragraph_linespace(paragraph)

def set_run_style(new_run, copy_run):
    new_run.font.italic= copy_run.font.italic
    new_run.font.strike= copy_run.font.strike
    new_run.font.bold= copy_run.font.bold
    new_run.font.size= copy_run.font.size
    new_run.font.color.rgb= copy_run.font.color.rgb
    new_run.underline= copy_run.underline
    new_run.style= copy_run.style

    # 字体名称设置需要特殊处理
    new_run.font.name = '微软雅黑'
    r = new_run._element.rPr.rFonts
    r.set(qn('w:eastAsia'),'微软雅黑')

def set_paragraph_linespace(paragraph):
    if hasattr(paragraph, "paragraph_format"):
        space_before=paragraph.paragraph_format.space_before
        space_after=paragraph.paragraph_format.space_after
        line_spacing=paragraph.paragraph_format.line_spacing
        line_spacing_rule=paragraph.paragraph_format.line_spacing_rule
        if space_before!=None:
            paragraph.paragraph_format.space_before=space_before
        if space_after!=None:
            paragraph.paragraph_format.space_after=space_after
        if line_spacing!=None:
            paragraph.paragraph_format.line_spacing=line_spacing
        if line_spacing_rule!=None:
            paragraph.paragraph_format.line_spacing_rule=line_spacing_rule

def check_image(run):
    if run.element.find('.//w:drawing', namespaces=run.element.nsmap) is not None:
        return True
    return False

# 去除照片
def clear_image(paragraph):
    for run in paragraph.runs:
        if check_image(run):
            run.clear()

def replace_paragraph_text(paragraph, texts, text_count, onlyText, appendTo):
    text=paragraph.text
    if check_text(text) and len(texts)>0:
        item=texts.pop(0)
        trans_text=item.get('text',"")
        if appendTo:
            if len(paragraph.runs)>0:
                paragraph.runs[-1].add_break()
                paragraph.runs[-1].add_text(trans_text)
            elif len(paragraph.hyperlinks)>0:
                paragraph.hyperlinks[-1].runs[-1].add_break()
                paragraph.hyperlinks[-1].runs[-1].add_text(trans_text)
        else:
            replaced=False
            if len(paragraph.runs)>0:
                for index,run in enumerate(paragraph.runs):
                    if not check_image(run):
                        if  not replaced:
                            run.text=trans_text
                            replaced=True
                        else:
                            run.clear()
            elif len(paragraph.hyperlinks)>0:
                for hyperlink in paragraph.hyperlinks:
                    for index,run in enumerate(hyperlink.runs):
                        if not check_image(run):
                            if  not replaced:
                                run.text=trans_text
                                replaced=True
                            else:
                                run.clear()

        text_count+=item.get('count',0)
        set_paragraph_linespace(paragraph)
    if onlyText:
        clear_image(paragraph)
        
def read_comments_from_docx(docx_path, texts):
    comments = []
    with zipfile.ZipFile(docx_path, 'r') as docx:
        # 尝试读取批注文件
        if 'word/comments.xml' in docx.namelist():
            with docx.open('word/comments.xml') as comments_file:
                # 解析 XML
                tree = ET.parse(comments_file)
                root = tree.getroot()
                
                # 定义命名空间
                namespace = {'ns0': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                
                # 查找所有批注
                for comment in root.findall('ns0:comment', namespace):
                    comment_id = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id')
                    author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author')
                    date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date')
                    text = ''.join(t.text for p in comment.findall('.//ns0:p', namespace) for r in p.findall('.//ns0:r', namespace) for t in r.findall('.//ns0:t', namespace))
                    append_comment(text, comment_id, texts)

def modify_comment_in_docx(docx_path, texts):
    # 创建一个临时文件名，保留原始路径
    temp_docx_path = os.path.join(os.path.dirname(docx_path), 'temp_' + os.path.basename(docx_path))

    # 打开原始 docx 文件
    with zipfile.ZipFile(docx_path, 'r') as docx:
        # 创建一个新的 docx 文件
        with zipfile.ZipFile(temp_docx_path, 'w') as new_docx:
            for item in docx.infolist():
                # 读取每个文件
                with docx.open(item) as file:
                    if item.filename == 'word/comments.xml':
                        # 解析批注 XML
                        tree = ET.parse(file)
                        root = tree.getroot()
                        
                        # 定义命名空间
                        namespace = {'ns0': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        
                        # 查找并修改批注
                        for comment in root.findall('ns0:comment', namespace):
                            text = ''.join(t.text for p in comment.findall('.//ns0:p', namespace) for r in p.findall('.//ns0:r', namespace) for t in r.findall('.//ns0:t', namespace))
                            if check_text(text):
                                for newitem in texts:
                                    # text_count+=newitem.get('count',0)
                                    new_text=newitem.get('text',"")
                                    comment_id=newitem.get('comment_id',"")
                                    # print("new_text:",new_text)
                                    # print("comment_id:",comment_id)
                                    # print("origin_id:",comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id'))
                                    if comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id') == comment_id:
                                        
                                        # 清除现有段落
                                        for p in comment.findall('.//ns0:t', namespace):
                                                # 删除 ns0:t 元素
                                            # comment.remove(p)  # 删除 ns0:t 元素

                                            # # 创建新的 ns0:t 元素
                                            # new_text_elem = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                                            # new_text_elem.text = new_text  # 设置新的文本内容

                                            # # 将新的 ns0:t 元素添加到段落中
                                            # r = ET.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r')  # 创建新的 run 元素
                                            # r.append(new_text_elem)  # 将新的 ns0:t 添加到 run 中
                                            # p.append(r)  # 将 run 添加到段落中
                                            p.text=new_text
                        # 打印修改后的 XML 内容
                        modified_xml = ET.tostring(root, encoding='utf-8', xml_declaration=True).decode('utf-8')
                        # print(modified_xml)
                        # 将修改后的 XML 写入新的 docx 文件
                        new_docx.writestr(item.filename, modified_xml)
                    else:
                        # 其他文件直接写入新的 docx 文件
                        new_docx.writestr(item.filename, file.read())

    # print(temp_docx_path)
    # 替换原始文件
    os.replace(temp_docx_path, docx_path)


def append_ins(text, ins_id, texts):
    if check_text(text):
        texts.append({"text": text, "type": "ins", "ins_id": ins_id, "complete": False})


def read_insstd_from_docx(docx_path, texts):
    document_ins = []
    namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    namespace14='{http://schemas.microsoft.com/office/word/2010/wordml}'
    with zipfile.ZipFile(docx_path, 'r') as docx:
        # 尝试读取批注文件
        if 'word/document.xml' in docx.namelist():
            with docx.open('word/document.xml') as document_file:
                # 解析 XML
                tree = ET.parse(document_file)
                root = tree.getroot()
                for element in root.findall(namespace + 'body'):
                    for p in element.findall(namespace + 'p'):
                        for ins in p.findall(namespace + 'ins'):
                            ins_id = ins.get(namespace + 'id')
                            for r in ins.findall(namespace + 'r'):
                                for t in r.findall(namespace + 't'):
                                    append_ins(t.text, ins_id, texts)
                    for sdt in element.findall(namespace + 'sdt'):
                        for sdtContent in sdt.findall(namespace + 'sdtContent'):
                            for p in sdtContent.findall(namespace + 'p'):                                
                                sdt_id = p.get(namespace14 + 'paraId')
                                for r in p.findall(namespace + 'r'):
                                    for t in r.findall(namespace + 't'):
                                        append_sdt(t.text, sdt_id, texts)
                                for ins in p.findall(namespace + 'ins'):
                                    for r in ins.findall(namespace + 'r'):
                                        for t in r.findall(namespace + 't'):
                                            append_sdt(t.text, sdt_id, texts)
                            


def append_sdt(text, sdt_id, texts):
    if check_text(text):
        texts.append({"text": text, "type": "sdt", "sdt_id": sdt_id, "complete": False})



def modify_inssdt_in_docx(docx_path, texts):
    print(texts,docx_path)
    temp_docx_path = os.path.join(os.path.dirname(docx_path), 'temp_std_' + os.path.basename(docx_path))
    with zipfile.ZipFile(docx_path, 'r') as docx:
        with zipfile.ZipFile(temp_docx_path, 'w') as new_docx:
            for item in docx.infolist():
                with docx.open(item) as file:
                    if item.filename == 'word/document.xml':
                        tree = ET.parse(file)
                        root = tree.getroot()
                        namespace = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                        namespace14='{http://schemas.microsoft.com/office/word/2010/wordml}'
                        for body in root.findall(namespace + 'body'):
                            for sdt in body.findall(namespace + 'sdt'):
                                for sdtContent in sdt.findall(namespace + 'sdtContent'):
                                    for p in sdtContent.findall(namespace + 'p') :
                                        for r in p.findall(namespace + 'r'):
                                            for t in r.findall(namespace + 't'):                                               
                                                text = t.text
                                                if check_text(text):
                                                    for newitem in texts:
                                                        new_text = newitem.get('text', "")
                                                        sdt_id = newitem.get('sdt_id', "")
                                                        if p.get(namespace14 + 'paraId') == sdt_id:
                                                            t.text = new_text
                                        for ins in p.findall(namespace + 'ins'):
                                            for r in ins.findall(namespace + 'r'):
                                                for t in r.findall(namespace + 't'):
                                                    text = t.text
                                                    if check_text(text):
                                                        for newitem in texts:
                                                            new_text = newitem.get('text', "")
                                                            sdt_id = newitem.get('sdt_id', "")
                                                            if p.get(namespace14 + 'paraId') == sdt_id:
                                                                t.text = new_text

                            for p in body.findall(namespace + 'p'):
                                for ins in p.findall(namespace + 'ins'):
                                    for r in ins.findall(namespace + 'r'):
                                        for t in r.findall(namespace + 't'):
                                            text = t.text
                                            if check_text(text):
                                                for newitem in texts:
                                                    new_text = newitem.get('text', "")
                                                    ins_id = newitem.get('ins_id', "")
                                                    if ins.get(namespace + 'id') == ins_id:
                                                        t.text = new_text
                        modified_xml = ET.tostring(root, encoding='utf-8', xml_declaration=True).decode('utf-8')
                        new_docx.writestr(item.filename, modified_xml)
                    else:
                        new_docx.writestr(item.filename, file.read())
    os.replace(temp_docx_path, docx_path)