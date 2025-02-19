import threading
import fitz
import re
import translate
import common
import io
import os
import sys
import time
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
# import pdfkit
import subprocess
import base64
import docx2pdf
import pdf2docx
import word
import copy
import shutil
from urllib.parse import quote
from io import BytesIO
from PIL import Image,ImageDraw
import pytesseract
import cv2
import numpy as np
import uuid
from pdfdeal import Doc2X
# from weasyprint import HTML

pytesseract.pytesseract.tesseract_cmd = r'/usr/local/bin/tesseract'

def start(trans):
    texts=[]
    src_pdf = fitz.open(trans['file_path'])
    # print(is_scan_pdf(src_pdf))
    # exit()
    # if is_scan_pdf(src_pdf):
    start_time = datetime.datetime.now()
    origin_docx_path=os.path.dirname(trans['file_path'])+"/"+trans['uuid']+".docx"
    target_docx_path=os.path.dirname(trans['file_path'])+"/"+trans['uuid']+"-translated.docx"
    target_pdf_path=os.path.dirname(trans['file_path'])+"/"+trans['uuid']+".pdf"
    # target_docx_path=re.sub(r"\.pdf",".docx",trans['target_file'], flags=re.I)
    # pdf_path=re.sub(r"\.pdf",".docx",trans['file_path'], flags=re.I)
    # print(target_pdf_path+"\n")
    # print(trans['storage_path']+"\n")
    # print(trans['target_file']+"\n")
    # print(os.path.join(trans['storage_path'], trans['target_filepath'])+"\n")
    pdftodocx(trans['file_path'], origin_docx_path)
    word_trans=copy.copy(trans)
    word_trans['file_path']=origin_docx_path
    word_trans['target_file']=target_docx_path
    word_trans['run_complete']=False;
    word_trans['extension']='.docx'
    text_count=0
    
    if word.start(word_trans):
        # print("word done")
        docxtopdf(target_docx_path, target_pdf_path)
        shutil.move(target_pdf_path, trans['target_file'])
        end_time = datetime.datetime.now()
        spend_time=common.display_spend(start_time, end_time)
        translate.complete(trans,text_count,spend_time)
        return True
    return False

    uuid=trans['uuid']
    html_path=trans['storage_path']+'/uploads/'+uuid
    trans['html_path']=html_path
    # read_pdf_html(trans['file_path'], html_path)
    # print(trans['storage_path']+'/uploads/pdf.html')
    # exit()
    # 允许的最大线程
    # print(trans)
    # wkhtmltopdf_bin=common.find_command_location("wkhtmltopdf")
    threads=trans['threads']
    if threads is None or int(threads)<0:
        max_threads=10
    else:
        max_threads=int(threads)
    # 当前执行的索引位置
    run_index=0
    start_time = datetime.datetime.now()
    # print(f'Source pdf file: {} \n', trans['file_path'])
    
    read_page_images(src_pdf, texts)
    
    text_count=0
    # translate.get_models()
    # exit()
    # read_page_html(src_pdf, texts, trans)
    # read_pdf_html(src_pdf, texts, trans)
    pdftohtml(trans['file_path'], html_path, texts)
    src_pdf.close()

    # print(texts)
    # exit()

    max_run=max_threads if len(texts)>max_threads else len(texts)
    event=threading.Event()
    before_active_count=threading.activeCount()
    while run_index<=len(texts)-1:
        if threading.activeCount()<max_run+before_active_count:
            if not event.is_set():
                # print("run_index:",run_index)
                thread = threading.Thread(target=translate.get,args=(trans,event,texts,run_index))
                thread.start()
                run_index+=1
            else:
                return False
    
    while True:
        if event.is_set():
            return False
        complete=True
        for text in texts:
            if not text['complete']:
                complete=False
        if complete:
            break
        else:
            time.sleep(1)


    # print(texts)

    write_to_html_file(html_path, texts)
    # config = pdfkit.configuration(wkhtmltopdf="/usr/local/bin/wkhtmltopdf")
    # with open(html_path) as f:
    #     pdfkit.from_file(f, trans['target_file'],options={"enable-local-file-access":True}, configuration=config)

    # print(trans['target_file'])

    end_time = datetime.datetime.now()
    spend_time=common.display_spend(start_time, end_time)
    translate.complete(trans,text_count,spend_time)
    return True


# def read_to_html(pages):

def read_page_html(pages, texts, trans):
    storage_path=trans['storage_path']
    uuid=trans['uuid']
    if is_scan_pdf(pages):
        for index,page in enumerate(pages):
            html=page.get_text("xhtml")
            images=re.findall(r"(data:image/\w+;base64,[^\"]+)", html)
            for i,image in enumerate(images):
                append_text(image, 'image', texts)

    else:
        for index,page in enumerate(pages):
            html=page.get_text("xhtml")
            # images=re.findall(r"(data:image/\w+;base64,[^\"]+)", html)
            # for i,image in enumerate(images):
            append_text(html,'text', texts)

def read_page_images(pages, texts):
    for index,page in enumerate(pages):
        html=page.get_text("xhtml")
        images=re.findall(r"(data:image/\w+;base64,[^\"]+)", html)
        for i,image in enumerate(images):
            append_text(image, 'image', texts)

def write_to_html_file(html_path,texts):
    with open(html_path, 'w+') as f:
        f.write('<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Strict//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd"><html><head><meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1"></head><body>')
        for item in texts:
            f.write(item.get("text", ""))
        f.write('</body></html>')
        f.close()

def read_block_text(pages,texts):
    text=""
    for page in pages:
        last_x0=0
        last_x1=0
        html=page.get_text("html")
        with open("test.html",'a+') as f:
            f.write(html)
            f.close()
        exit()
        for block in page.get_text("blocks"):
            current_x1=block[2]
            current_x0=block[0]
            # 对于每个文本块，分行并读取
            if block[5]==0 or abs(current_x1-last_x1)>12 or abs(current_x0-last_x0)>12:
                append_text(text, "text", texts)
                text=block[4].replace("\n","")
            else:
                text=text+(block[4].replace("\n",""))
            last_x1=block[2]
            last_x0=block[0]
    append_text(text, "text", texts)

def write_block_text(pages,newpdf,texts):
    text=""
    for page in pages:
        last_x0=0
        last_x1=0
        last_y0=0
        new_page = newpdf.new_page(width=page.rect.width, height=page.rect.height)
        font=fitz.Font("helv")
        for block in page.get_text("blocks"):
            current_x1=block[2]
            current_x0=block[0]
            current_y0=block[1]
            # 对于每个文本块，分行并读取
            if block[5]==0 or abs(current_x1-last_x1)>12 or abs(current_x0-last_x0)>12 and len(texts)>0:
                item=texts.pop(0)
                trans_text=item.get("text","")
                new_page.insert_text((last_x0,last_y0), trans_text, fontsize=12,fontname="Helvetica", overlay=False)
                text=block[4].replace("\n","")
            else:
                text=text+(block[4].replace("\n",""))
            last_x1=block[2]
            last_x0=block[0]
            last_y0=block[1]
    if check_text(text) and len(texts):
        new_page.insert_text((last_x0,last_y0), trans_text, fontsize=12, overlay=False)

def write_page_text(pages,newpdf,texts):
    for page in pages:
        text=page.get_text("text")
        new_page = newpdf.new_page(width=page.rect.width, height=page.rect.height)
        if check_text(text) and len(texts)>0:
            item=texts.pop(0)
            text=item.get("text","")
            new_page.insert_text((0,0), text, fontsize=12, overlay=False)

def read_row(pages,texts):
    text=""
    for page in pages:
        # 获取页面的文本块
        for block in page.get_text("blocks"):
            # 对于每个文本块，分行并读取
            if block[5]==0:
                append_text(text, 'text', texts)
                text=block[4]
            else:
                text=text+block[4]

def write_row(newpdf, texts, page_width, page_height):
    text_count=0
    new_page = newpdf.new_page(width=page_width, height=page_height)
    for text in texts:
        print(text['text'])
        # draw_text_avoid_overlap(new_page, text['text'],text['block'][0],text['block'][1], 16)
        new_page.insert_text((text['block'][0],text['block'][1]),text['text'], fontsize=16)
        return



def append_text(text, content_type, texts):
    if check_text(text):
        # print(text)
        texts.append({"text":text,"type":content_type, "complete":False})


def check_text(text):
    return text!=None and len(text)>0 and not common.is_all_punc(text) 

def draw_text_avoid_overlap(page, text, x, y, font_size):
    """
    在指定位置绘制文本，避免与现有文本重叠。
    """
    text_length = len(text) * font_size  # 估算文本长度
    while True:
        text_box = page.get_textbox((x, y, x + text_length, y + font_size))
        if not text_box:
            break  # 没有重叠的文本，退出循环
        y += font_size + 1  # 移动到下一个位置
 
    page.insert_text((x,y),text, fontsize=font_size)


def draw_table(page, table_data, x, y, width, cell_height):
    # 表格的列数
    cols = len(table_data[0])
    rows = len(table_data)
    
    # 绘制表格
    for i in range(rows):
        for j in range(cols):
            # 文字写入
            txt = table_data[i][j]
            page.insert_text((x, y), txt)
            # 绘制单元格边框 (仅边界线)
            # 左边
            page.draw_line((x, y),( x+width/cols, y), width=0.5)
            # 上边
            if i == 0:
                page.draw_line((x, y), (x, y+cell_height), width=0.5)
            # 右边
            if j == cols-1:
                page.draw_line((x+width/cols, y), (x+width/cols, y+cell_height), width=0.5)
            # 下边
            if i == rows-1:
                page.draw_line((x, y+cell_height), (x+width/cols, y+cell_height), width=0.5)
            # 移动到下一个单元格
            x += width/cols
        # 移动到下一行
        x = 0
        y += cell_height

def wrap_text(text, width):
    words = text.split(' ')
    lines = []
    line = ""
    for word in words:
        if len(line.split(' ')) >= width:
            lines.append(line)
            line = ""
        if len(line + word + ' ') <= width * len(word):
            line += word + ' '
        else:
            lines.append(line)
            line = word + ' '
    if line:
        lines.append(line)
    return lines


def is_paragraph(block):
    # 假设一个段落至少有两行
    if len(block) < 2:
        return False
    # 假设一个段落的行间隔较大
    if max([line.height for line in block]) / min([line.height for line in block]) > 1.5:
        return True
    return False

def is_next_line_continuation(page, current_line, next_line_index):
    # 判断下一行是否是当前行的继续
    return abs(next_line_index - current_line) < 0.1

def print_texts(texts):
    for item in texts:
        print(item.get("text"))

def is_scan_pdf(pages):
     for index,page in enumerate(pages):
        html=page.get_text("xhtml")
        images=re.findall(r"(data:image/\w+;base64,[^\"]+)", html)
        text=page.get_text()
        print(images)
        print(text)
        if text=="" and len(images)>0:
            return True
        else:
            return False

def read_pdf_html(pages, texts, trans):
    for index,page in enumerate(pages):
        target_html="{}-{}.html".format(trans['html_path'], page_index)
        if os.path.exists(target_html):
            os.remove(target_html)
        dftohtml_path = shutil.which("pdftohtml")
        if pdftohtml_path is None:
            raise Exception("未安装pdftohtml")
        subprocess.run([dftohtml_path,"-c","-l", page_index, trans['file_path'], trans['html_path']])
        if not os.path.exists(target_html):
            raise Exception("无法生成html")
        # append_text(html,'text', texts)


def pdftohtml(pdf_path, html_path,texts):
    target_html="{}-html.html".format(html_path)
    if os.path.exists(target_html):
        os.remove(target_html)
    pdftohtml_path = shutil.which("pdftohtml")
    if pdftohtml_path is None:
        raise Exception("未安装pdftohtml")
    subprocess.run([pdftohtml_path,"-c","-s", pdf_path, html_path])
    if not os.path.exists(target_html):
        raise Exception("无法生成html")
    with open(target_html, 'r') as f:
        content=f.read()
        print(content)
        append_text(content, 'text', texts)


def pdftodocx(pdf_path, docx_path):
    print(docx_path)
    if os.path.exists(docx_path):
        os.remove(docx_path)
    print(pdf_path)
    try:
        cv = pdf2docx.Converter(pdf_path)
        cv.debug_page(0)
        cv.convert(docx_path, start=0,end=1,multi_processing=False)
        cv.close()
        #exit()
    except Exception as e:
        print("error")
        pdf2docxNext(pdf_path, docx_path)

def pdf2docxNext(pdf_path, docx_path):
    try:
        # 创建一个新的 DOCX 文档
        doc = Document()
        # 打开 PDF 文件
        pdf_document = fitz.open(pdf_path)
        # 遍历 PDF 的每一页
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            fonts=page.get_fonts()
            # 提取文本
            # 提取文本和样式信息
            text_dict = page.get_text("dict")
            
            # 遍历文本块
            for block in text_dict["blocks"]:
                if block["type"] == 0:  # 只处理文本块
                    for line in block["lines"]:
                        for span in line["spans"]:
                            text = span["text"]
                            font_size = span["size"]  # 字体大小
                            font_color = span["color"]  # 字体颜色
                            
                            # 创建段落
                            paragraph = doc.add_paragraph()
                            run = paragraph.add_run(text)
                            
                            # 设置字体大小
                            run.font.size = Pt(font_size)
                            
                            # 设置字体颜色
                            if font_color:
                                run.font.color.rgb = RGBColor(
                                    (font_color >> 16) & 0xFF,  # R
                                    (font_color >> 8) & 0xFF,   # G
                                    font_color & 0xFF            # B
                                )
                elif block["type"] == 1:
                    # 提取图像
                    try:
                        img_index = block["image"]
                        base_image = pdf_document.extract_image(img_index)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        # 将图像添加到 DOCX
                        image_stream = BytesIO(image_bytes)
                        doc.add_picture(image_stream, width=None)  # 可以指定宽度
                    except Exception as e:
                        print("图片无法解析")
                    


            # 添加分页符
            doc.add_page_break()
        
        # 保存 DOCX 文件
        doc.save(docx_path)
        pdf_document.close()
    except Exception as e:
        raise("pdf转docx失败")

def docxtopdf(docx_path, pdf_path):
    if os.path.exists(pdf_path):
        os.remove(pdf_path)
    sys.path.append("/usr/local/bin")
    unoconv_path = shutil.which("unoconv")
    if unoconv_path is None:
        raise Exception("未安装unoconv")
    target_path_dir=os.path.dirname(pdf_path)
    if not os.path.exists(target_path_dir):
        os.makedirs(target_path_dir, mode=0o777, exist_ok=True)
    # target_pdf = fitz.Document()
    # target_pdf.new_page()
    # target_pdf.save(pdf_path)
    # target_pdf.close()
    # subprocess.run([unoconv_path,"-f","pdf","-e","UTF-8","-o",target_path_dir, docx_path])
    # subprocess.run([unoconv_path,"-f","pdf","-e","UTF-8","-o",target_path_dir, docx_path])
    print("{} -f pdf -o {} {}".format(unoconv_path,pdf_path, docx_path))
    # subprocess.run("{} -f pdf -o {} {}".format(unoconv_path, pdf_path, docx_path), shell=True)
    command = [unoconv_path, "-f", "pdf", "-o", pdf_path, docx_path]
    subprocess.run(command)
    print("done")

def create_temp_file(suffix='.png'):
    temp_dir = '/tmp'  # 或者使用其他临时目录
    filename = f"{uuid.uuid4()}{suffix}"
    return os.path.join(temp_dir, filename)

def pdf_to_text_with_ocr(pdf_path, docx_path, origin_lang):
    # if not is_tesseract_installed():
    #     raise Exception("Tesseract未安装,无法进行OCR")
    
    document = fitz.open(pdf_path)
    docx = Document()

    for page_num in range(len(document)):
        page = document.load_page(page_num)
        pix = page.get_pixmap()
        img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
        
        # 转换为灰度图像
        img = img.convert('L')
        
        # 将图像保存到内存中的字节流
        img_byte_arr = io.BytesIO()
        img.save(img_byte_arr, format='PNG')
        img_byte_arr = img_byte_arr.getvalue()
        
        try:
            # 使用 Tesseract 命令行工具
            process = subprocess.Popen(
                ['/usr/local/bin/tesseract', 'stdin', 'stdout', '-l', origin_lang, '--oem', '3', '--psm', '6'],
                stdin=subprocess.PIPE,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            stdout, stderr = process.communicate(input=img_byte_arr)
            
            if process.returncode != 0:
                raise subprocess.CalledProcessError(process.returncode, process.args, stdout, stderr)
            
            text = stdout.decode('utf-8').strip()
            
            # 移除空行和多余的空格
            text = '\n'.join(line.strip() for line in text.splitlines() if line.strip())
            
        except subprocess.CalledProcessError as e:
            print(f"OCR处理页面 {page_num + 1} 时出错: {str(e)}")
            text = ""  # 如果出错，使用空字符串
        
        paragraph = docx.add_paragraph()
        run = paragraph.add_run(text)
        run.font.size = Pt(12)

    document.close()
    docx.save(docx_path)

def is_scanned_pdf(pdf_path):
    document = fitz.open(pdf_path)
    
    # 只检查前几页，通常足以判断
    pages_to_check = min(5, len(document))
    
    for page_num in range(pages_to_check):
        page = document[page_num]
        
        # 检查文本
        text = page.get_text().strip()
        if text:
            document.close()
            return False
        
        # 检查图像
        image_list = page.get_images()
        if len(image_list) > 0:
            # 如果页面只包含一个大图像，很可能是扫描件
            if len(image_list) == 1:
                xref = image_list[0][0]
                img = document.extract_image(xref)
                if img:
                    pix = fitz.Pixmap(img["image"])
                    # 如果图像覆盖了大部分页面，可能是扫描件
                    if pix.width > page.rect.width * 0.9 and pix.height > page.rect.height * 0.9:
                        document.close()
                        return True
    
    document.close()
    return True  # 如果没有找到文本，默认认为是扫描件

def is_tesseract_installed():
    tesseract_path = "/usr/local/bin/tesseract"
    return os.path.isfile(tesseract_path) and os.access(tesseract_path, os.X_OK)

def use_doc2x_revert_pdf_to_docx(dox2x_api_key, pdf_file, docx_path):
    client = Doc2X(apikey=dox2x_api_key,debug=False)
    success, failed, flag = client.pdf2file(
        pdf_file=pdf_file,
        output_path=docx_path,
        output_format="docx",
    )
    if len(success)>0 and success[0]!="":
        return (True,success[0])
    else:
        return (False,failed[0]["error"])

# def save_image(base64_data, path):
#     image_data = base64.b64decode(base64_data)
#     # 将字节数据写入内存中的文件对象
#     image_file = BytesIO(image_data)
#     # 从内存中的文件对象创建Image对象
#     image = Image.open(image_file)
#     # 保存图片到文件系统
#     image.sav/e(path)

